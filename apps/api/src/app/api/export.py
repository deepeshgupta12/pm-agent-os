from __future__ import annotations

from fastapi import APIRouter, Depends, HTTPException
from fastapi.responses import Response
from sqlalchemy.orm import Session

from app.api.deps import require_user, require_workspace_access
from app.db.session import get_db
from app.db.models import Artifact, Run, User
from app.core.pdf_export import markdown_to_pdf_bytes

router = APIRouter(tags=["export"])


def _ensure_artifact_access(db: Session, art: Artifact, user: User) -> None:
    run = db.get(Run, art.run_id)
    if not run:
        raise HTTPException(status_code=404, detail="Artifact not found")

    # viewer+ can export
    require_workspace_access(str(run.workspace_id), db, user)


@router.get("/artifacts/{artifact_id}/export/pdf")
def export_artifact_pdf(
    artifact_id: str,
    db: Session = Depends(get_db),
    user: User = Depends(require_user),
):
    art = db.get(Artifact, artifact_id)
    if not art:
        raise HTTPException(status_code=404, detail="Artifact not found")

    _ensure_artifact_access(db, art, user)

    pdf_bytes = markdown_to_pdf_bytes(title=art.title, markdown=art.content_md)
    filename = f"{art.logical_key}-v{art.version}.pdf".replace(" ", "_")

    return Response(
        content=pdf_bytes,
        media_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


@router.get("/artifacts/{artifact_id}/export/docx")
def export_artifact_docx(
    artifact_id: str,
    db: Session = Depends(get_db),
    user: User = Depends(require_user),
):
    """
    V1 simple Markdown -> DOCX:
    - # / ## / ### => headings
    - "- " => bullet points
    - other lines => paragraphs
    """
    try:
        from docx import Document as DocxDocument
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"python-docx missing: {e}")

    art = db.get(Artifact, artifact_id)
    if not art:
        raise HTTPException(status_code=404, detail="Artifact not found")

    _ensure_artifact_access(db, art, user)

    doc = DocxDocument()
    doc.add_heading(art.title or "Artifact", level=1)

    lines = (art.content_md or "").splitlines()
    for line in lines:
        s = line.rstrip()

        if s.startswith("### "):
            doc.add_heading(s[4:].strip(), level=3)
            continue
        if s.startswith("## "):
            doc.add_heading(s[3:].strip(), level=2)
            continue
        if s.startswith("# "):
            doc.add_heading(s[2:].strip(), level=1)
            continue

        if s.startswith("- "):
            doc.add_paragraph(s[2:].strip(), style="List Bullet")
            continue

        if not s.strip():
            doc.add_paragraph("")
            continue

        doc.add_paragraph(s)

    from io import BytesIO

    buf = BytesIO()
    doc.save(buf)
    docx_bytes = buf.getvalue()

    filename = f"{art.logical_key}-v{art.version}.docx".replace(" ", "_")

    return Response(
        content=docx_bytes,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )