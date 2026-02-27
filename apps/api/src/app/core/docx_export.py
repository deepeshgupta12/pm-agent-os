from __future__ import annotations

import re
from io import BytesIO
from typing import Optional, Tuple

try:
    from docx import Document as DocxDocument
    from docx.shared import Pt
except Exception:  # pragma: no cover
    DocxDocument = None  # type: ignore
    Pt = None  # type: ignore


_HEADING_RE = re.compile(r"^(#{1,6})\s+(.*)$")
_BULLET_RE = re.compile(r"^(\-|\*)\s+(.*)$")
_NUMBERED_RE = re.compile(r"^(\d+)\.\s+(.*)$")
_LINK_MD_RE = re.compile(r"\[([^\]]+)\]\(([^)]+)\)")


def _split_link_markdown(text: str) -> list[Tuple[str, Optional[str]]]:
    """
    Splits markdown links into runs: [(text, None), (link_text, url), ...]
    Keeps it simple (no nested/parens edge cases).
    """
    out: list[Tuple[str, Optional[str]]] = []
    i = 0
    for m in _LINK_MD_RE.finditer(text):
        if m.start() > i:
            out.append((text[i : m.start()], None))
        out.append((m.group(1), m.group(2)))
        i = m.end()
    if i < len(text):
        out.append((text[i:], None))
    return out


def markdown_to_docx_bytes(title: str, markdown: str) -> bytes:
    """
    Markdown -> DOCX (improved, still intentionally light-weight)
    Supports:
      - #..###### headings
      - bullets (-, *)
      - numbered lists (1. 2. ...)
      - fenced code blocks ``` ... ```
      - blank lines
      - inline links [text](url) (stored as text + url in parentheses; reliable & simple)
    """
    if DocxDocument is None:
        raise RuntimeError("python-docx is not installed")

    doc = DocxDocument()

    # Title
    doc.add_heading(title or "Artifact", level=1)

    lines = (markdown or "").splitlines()
    in_code = False
    code_buf: list[str] = []

    def flush_code():
        nonlocal code_buf
        if not code_buf:
            return
        p = doc.add_paragraph()
        run = p.add_run("\n".join(code_buf))
        # monospace-ish
        try:
            run.font.name = "Courier New"
            run.font.size = Pt(9)  # type: ignore[misc]
        except Exception:
            pass
        code_buf = []

    for raw in lines:
        s = raw.rstrip("\n")

        # fenced code
        if s.strip().startswith("```"):
            if in_code:
                # close
                in_code = False
                flush_code()
            else:
                # open
                in_code = True
                code_buf = []
            continue

        if in_code:
            code_buf.append(s)
            continue

        # headings
        hm = _HEADING_RE.match(s.strip())
        if hm:
            level = min(len(hm.group(1)), 6)
            text = hm.group(2).strip()
            # python-docx supports 1..9; we map 1..6 -> 1..6
            doc.add_heading(text, level=level)
            continue

        # bullets
        bm = _BULLET_RE.match(s)
        if bm:
            txt = bm.group(2).strip()
            doc.add_paragraph(txt, style="List Bullet")
            continue

        # numbered
        nm = _NUMBERED_RE.match(s)
        if nm:
            txt = nm.group(2).strip()
            # Default Word numbering style name varies; "List Number" usually exists
            try:
                doc.add_paragraph(txt, style="List Number")
            except Exception:
                doc.add_paragraph(f"{nm.group(1)}. {txt}")
            continue

        # blank
        if not s.strip():
            doc.add_paragraph("")
            continue

        # paragraph with inline links (simple)
        p = doc.add_paragraph()
        parts = _split_link_markdown(s)
        for text_part, url in parts:
            if not text_part:
                continue
            if url:
                # reliable: write "text (url)" instead of true hyperlink relationship
                p.add_run(text_part)
                p.add_run(f" ({url})")
            else:
                p.add_run(text_part)

    # flush if file ended mid-code
    if in_code:
        flush_code()

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()